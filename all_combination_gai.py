import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn
import pandas as pd
import xlrd

def gen_docfile(data_path, doc_file_path):
    # 建一个文档
    document = Document()
    style = document.styles['Normal']
    document.styles['Normal'].font.name = u'Times New Roman'  # 设置西文字体
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')  # 设置中文字体使用宋体
    document.styles['Normal'].font.size = Pt(14)

    # 报告大标题
    t0 = document.add_heading('')
    t0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    t0.add_run().bold = True
    t0.add_run('高端产业功能区月度监测报告').font.size = Pt(18)

    # 读取数据
    df_1 = pd.read_excel(data_path + '\企业动态.xlsx')
    df_2 = pd.read_excel(data_path + '\月新增注册.xlsx')
    df_3 = pd.read_excel(data_path + '\投融资.xlsx')
    df_4 = pd.read_excel(data_path + '\创新发展.xlsx')

    # 添加一个段落
    p1 = document.add_paragraph()
    text = u'——%s'\
        % (str(df_1['功能区'][i]))
    r = p1.add_run(text)
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r.font.color.rgb = RGBColor(255, 0, 0)

    # 首行缩进两个字符
    paragraph_format = style.paragraph_format
    paragraph_format.first_line_indent = Cm(0.74)
    paragraph_format.line_spacing = 1.5  # 修改行间距
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # 第一段
    t1 = document.add_heading('', level=1)
    t1.add_run('一、企业动态（截止到2022年5月底最新现状数据）').bold = True
    text = u'截至2022年5月31日，%s有企业%s家，注册资本%s亿元。' \
           u'其中，国有企业%s家，注册资本总额%s亿元；外商投资企业%s家，注册资本总额%s亿元；' \
           u'高新技术企业%s家，世界500强企业%s家，上市企业%s家（为港股企业），瞪羚企业%s家，独角兽企业%s家。'\
    % (str(df_1['功能区'][i]),
      str(df_1['企业数量'][i]),
      str(round(df_1['注册资本总额（亿）'][i], 2)),
      str(df_1['国有企业数量'][i]),
      str(round(df_1['国有企业注册金额（亿）'][i], 2)),
      str(df_1['外商投资企业数量'][i]),
      str(round(df_1['外商投资企业注册资本总额（亿）'][i], 2)),
      str(df_1['高新技术企业'][i]),
      str(df_1['世界500强'][i]),
      str(df_1['港股'][i]),
      str(df_1['瞪羚企业'][i]),
      str(df_1['独角兽企业'][i]))
    document.add_paragraph(text)

    text = u'5月新增企业%s家，新增注册资本%s亿元，注销企业%s家。'\
        % (str(df_2['新增注册'][i]),
        str(round(df_2['新增注册资本总额（亿）'][i], 2)),
        str(df_2['新增注销'][i])
      )
    document.add_paragraph(text)

    # #插入表格
    table_name_1 = document.add_paragraph(u'表1：%s5月新增注册企业情况'\
                                          % (str(df_1['功能区'][i])))
    table_name_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if os.path.exists(data_path + '\当月新增-表格\新增注册-' + df_1['功能区'][i] + '.xlsx'):
        wb_new_registered = xlrd.open_workbook(data_path + '\当月新增-表格\新增注册-' + df_1['功能区'][i] + '.xlsx')
        sheetname_new_registered = wb_new_registered.sheet_names()[0]
        sheet_c_new_registered = wb_new_registered.sheet_by_index(0)

        table = document.add_table(rows=sheet_c_new_registered.nrows, cols=4, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'企业名称'
        hdr_cells[1].text = u'企业法人'
        hdr_cells[2].text = u'注册资本（万元）'
        hdr_cells[3].text = u'企业简介'

        for j in range(1, sheet_c_new_registered.nrows):
            row_data = sheet_c_new_registered.row_values(j)
            #     print(row_data[0])
            table.cell(j, 0).text = str(row_data[0])
            table.cell(j, 1).text = str(row_data[1])
            table.cell(j, 2).text = str(row_data[2])
            table.cell(j, 3).text = str(row_data[3])
    else:
        print(df_1['功能区'][i] + ' 没有新增注册的企业')

    # #插入表格
    table_name_2 = document.add_paragraph(u'表2：%s5月新增注销企业情况' \
                                          % (str(df_1['功能区'][i])))
    table_name_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if os.path.exists(data_path + '\当月新增-表格\新增注销-' + df_1['功能区'][i] + '.xlsx'):
        wb_new_out = xlrd.open_workbook(data_path + '\当月新增-表格\新增注销-' + df_1['功能区'][i] + '.xlsx')
        sheetname_new_out = wb_new_out.sheet_names()[0]
        sheet_c_new_out = wb_new_out.sheet_by_index(0)

        table = document.add_table(rows=sheet_c_new_out.nrows, cols=4, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'企业名称'
        hdr_cells[1].text = u'企业法人'
        hdr_cells[2].text = u'注册资本（万元）'
        hdr_cells[3].text = u'企业简介'

        for j in range(1, sheet_c_new_out.nrows):
            row_data = sheet_c_new_out.row_values(j)
            #     print(row_data[0])
            table.cell(j, 0).text = str(row_data[0])
            table.cell(j, 1).text = str(row_data[1])
            table.cell(j, 2).text = str(row_data[2])
            table.cell(j, 3).text = str(row_data[3])
    else:
        print(df_1['功能区'][i] + ' 没有新增注销的企业')

    # 第二段
    document.add_heading(u'二、投融资情况', level=1)
    text = u'投融资方面，5月，%s企业获得投资%s元，' \
         u'其中，境外资本投资%s元；%s内企业对外投资%s元，主要投向%s等地，主要集中在%s产业领域。' \
           % (str(df_3['功能区'][i]),
              str(df_3['融资金额（元）'][i]),
              str(df_3['境外投资总金额（元）'][i]),
              str(df_3['功能区'][i]),
              str(df_3['对外投资总金额（元）'][i]),
              str(df_3['对外投资地区'][i]),
              str(df_3['对外投资行业分布'][i]))
    document.add_paragraph(text)

    # #插入表格
    table_name_3 = document.add_paragraph(u'表3：%s企业融资情况明细' \
                                          % (str(df_3['功能区'][i])))
    table_name_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if os.path.exists(data_path + '\投融资-表格\融资情况-' + df_1['功能区'][i] + '.xlsx'):
        wb_money_out = xlrd.open_workbook(data_path + '\投融资-表格\融资情况-' + df_1['功能区'][i] + '.xlsx')
        sheetname_money_out = wb_money_out.sheet_names()[0]
        sheet_c_money_out = wb_money_out.sheet_by_index(0)

        table = document.add_table(rows=sheet_c_money_out.nrows, cols=6, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'日期'
        hdr_cells[1].text = u'企业名称'
        hdr_cells[2].text = u'法人名称'
        hdr_cells[3].text = u'投资人'
        hdr_cells[4].text = u'融资金额（万元）'
        hdr_cells[5].text = u'融资轮次'

        for j in range(1, sheet_c_money_out.nrows):
            row_data = sheet_c_money_out.row_values(j)
            #     print(row_data[0])
            table.cell(j, 0).text = str(row_data[0])
            table.cell(j, 1).text = str(row_data[1])
            table.cell(j, 2).text = str(row_data[2])
            table.cell(j, 3).text = str(row_data[3])
            table.cell(j, 4).text = str(row_data[4])
            table.cell(j, 5).text = str(row_data[5])
    else:
        print(df_1['功能区'][i] + ' 没有融资信息')

    # #插入表格
    table_name_4 = document.add_paragraph(u'表4：%s企业对外投资明细'\
                                          % str(df_3['功能区'][i]))
    table_name_4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if os.path.exists(data_path + '\投融资-表格\对外投资-' + df_1['功能区'][i] + '.xlsx'):

        wb_money_in = xlrd.open_workbook(data_path + '\投融资-表格\对外投资-' + df_1['功能区'][i] + '.xlsx')
        sheetname_money_in = wb_money_in.sheet_names()[0]
        sheet_c_money_in = wb_money_in.sheet_by_index(0)

        table = document.add_table(rows=sheet_c_money_in.nrows, cols=6, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'日期'
        hdr_cells[1].text = u'企业名称'
        hdr_cells[2].text = u'法人名称'
        hdr_cells[3].text = u'被投资企业'
        hdr_cells[4].text = u'投资金额（万元）'
        hdr_cells[5].text = u'企业简介'

        for j in range(1, sheet_c_money_in.nrows):
            row_data = sheet_c_money_in.row_values(j)
            #     print(row_data[0])
            table.cell(j, 0).text = str(row_data[0])
            table.cell(j, 1).text = str(row_data[1])
            table.cell(j, 2).text = str(row_data[2])
            table.cell(j, 3).text = str(row_data[3])
            table.cell(j, 4).text = str(row_data[4])
            table.cell(j, 5).text = str(row_data[5])
    else:
        print(df_1['功能区'][i] + ' 没有对外投资')

    # 第三段
    document.add_heading(u'三、创新发展', level=1)
    text = u'截至2022年5月31日，%s功能区专利申请量为%s件。5月专利申请量新增%s件。'\
     % (str(df_4['功能区'][i]),
       str(df_4['专利申请量'][i]),
       str(df_4['专利申请当月新增量'][i])
       )
    document.add_paragraph(text)
    # #插入表格
    table_name_5 = document.add_paragraph(u'表5：%s授权专利情况'\
            % str(df_4['功能区'][i]))
    table_name_5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if os.path.exists(data_path + '\创新发展-表格\创新发展-' + df_1['功能区'][i] + '.xlsx'):
        wb_patent = xlrd.open_workbook(data_path + '\创新发展-表格\创新发展-' + df_1['功能区'][i] + '.xlsx')
        sheetname_patent = wb_patent.sheet_names()[0]
        sheet_c_patent = wb_patent.sheet_by_index(0)

        table = document.add_table(rows=sheet_c_patent.nrows, cols=4, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'专利名称'
        hdr_cells[1].text = u'申请机构/个人'
        hdr_cells[2].text = u'所属行业'
        hdr_cells[3].text = u'申请时间'

        for j in range(1, sheet_c_patent.nrows):
            row_data = sheet_c_patent.row_values(j)
            #     print(row_data[0])
            table.cell(j, 0).text = str(row_data[0])
            table.cell(j, 1).text = str(row_data[1])
            table.cell(j, 2).text = str(row_data[2])
            table.cell(j, 3).text = str(row_data[3])
    else:
        print(df_1['功能区'][i] + ' 没有创新发展')

    # 第四段
    document.add_heading(u'四、重点项目', level=1)
    text = u'请手动补充。'
    document.add_paragraph(text)

    # 第五段
    document.add_heading(u'五、政策动态', level=1)
    # #插入表格
    table_name_6 = document.add_paragraph(u'表6：2022年5月%s相关政策动态' \
                                          % str(df_1['功能区'][i]))
    table_name_6.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if os.path.exists(data_path + '\政策-表格\政策信息-' + df_1['功能区'][i] + '.xlsx'):
        wb_policy = xlrd.open_workbook(data_path + '\政策-表格\政策信息-' + df_1['功能区'][i] + '.xlsx')
        sheetname_policy = wb_policy.sheet_names()[0]
        sheet_c_policy = wb_policy.sheet_by_index(0)

        table = document.add_table(rows=sheet_c_policy.nrows, cols=3, style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'政策标题'
        hdr_cells[1].text = u'政策内容'
        hdr_cells[2].text = u'发布时间'

        for j in range(1, sheet_c_policy.nrows):
            row_data = sheet_c_policy.row_values(j)
            #     print(row_data[0])
            table.cell(j, 0).text = str(row_data[0])
            table.cell(j, 1).text = str(row_data[1])
            table.cell(j, 2).text = str(row_data[2])
    else:
        print(df_1['功能区'][i] + ' 没有政策动态')

    document.save(doc_file_path + df_1['功能区'][i] + '.docx')


for i in range(0, 12):
    gen_docfile(r'.\data_5', r'.\report_5\\')